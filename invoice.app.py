import os
import re
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import tempfile
import shutil
from datetime import datetime
import base64
import zipfile
import subprocess
from typing import Optional
from export_excel import export_parking_list

# Configuration
TEMPLATE_PATH = "invoice_template.docx"
OUTPUT_FOLDER = "generated_invoices"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def sanitize_filename(name: str) -> str:
    """Make strings safe for filenames"""
    return re.sub(r'[\\/:*?"<>|]', '_', name)

def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert DOCX to PDF using LibreOffice"""
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", 
             os.path.dirname(pdf_path), docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=30
        )
        
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        temp_pdf = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
        
        if os.path.exists(temp_pdf):
            os.rename(temp_pdf, pdf_path)
            return True
        return False
        
    except Exception as e:
        st.error(f"PDF conversion failed. Error: {str(e)}")
        return False

def validate_pdf(pdf_path: str) -> bool:
    """Verify PDF is valid"""
    try:
        with open(pdf_path, "rb") as f:
            return f.read(4) == b"%PDF"
    except:
        return False

def generate_pdf_from_template(
    template_path: str,
    row_data: dict,
    output_folder: str,
    invoice_number: int
) -> Optional[str]:
    """Generate PDF invoice from template"""
    try:
        doc = Document(template_path)
        style = doc.styles['Normal']
        style.font.size = Pt(8)

        # Ensure all financial values are properly formatted
        financial_fields = [
            'PER CHARGES', 'PARKING CHARGES', 'TOTAL CHARGES', 
            'TOTAL CHARGES_SUM', 'RATE', 'CALCULATED_CHARGES'
        ]
        
        # Format values
        formatted_data = {
            k: f"{float(v):.2f}" if k in financial_fields and str(v).replace('.','',1).isdigit() else str(v)
            for k, v in row_data.items()
        }
        
        current_date = datetime.now().strftime("%Y-%m-%d")
        formatted_data.update({
            "DATE": current_date,
            "INVOICE NUMBER": str(invoice_number),
            "TRACKING NUMBER": row_data.get("TRACKING NUMBER", ""),
            "TERMS": row_data.get("TERMS", ""),
            "TOTAL QTY": row_data.get("TOTAL QTY", "")
        })

        # Add invoice header
        if len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0]
            first_para.text = f"Invoice #: {invoice_number}\nDate: {current_date}\n" + first_para.text
            for run in first_para.runs:
                run.font.size = Pt(10)
                run.bold = True

        # Replace placeholders - including the rate
        for paragraph in doc.paragraphs:
            for key, value in formatted_data.items():
                for ph in [f"{{{{{key}}}}}", f"{{{{{key}.}}}}"]:
                    if ph in paragraph.text:
                        paragraph.text = paragraph.text.replace(ph, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in formatted_data.items():
                        for ph in [f"{{{{{key}}}}}", f"{{{{{key}.}}}}"]:
                            if ph in cell.text:
                                cell.text = cell.text.replace(ph, value)

        # Generate filename
        customer = sanitize_filename(formatted_data.get("MARK", "Customer"))
        pdf_name = f"Invoice_{invoice_number}_{customer}.pdf"
        pdf_path = os.path.join(output_folder, pdf_name)
        
        # Save and convert
        temp_docx = os.path.join(output_folder, f"temp_{invoice_number}.docx")
        doc.save(temp_docx)
        
        if convert_docx_to_pdf(temp_docx, pdf_path) and validate_pdf(pdf_path):
            os.remove(temp_docx)
            return pdf_path
        
        os.remove(temp_docx)
        return None

    except Exception as e:
        st.error(f"Template processing failed: {str(e)}")
        return None

def update_notification_sheet(output_folder: str, pdf_name: str, customer: str, 
                            invoice_number: int, contact: str, total: str):
    """Update tracking spreadsheet"""
    sheet_path = os.path.join(output_folder, "notification_log.xlsx")
    new_data = pd.DataFrame([{
        "Customer": customer,
        "Invoice": invoice_number,
        "Contact": contact,
        "Amount": total,
        "File": pdf_name
    }])
    
    if os.path.exists(sheet_path):
        try:
            existing = pd.read_excel(sheet_path)
            updated = pd.concat([existing, new_data])
            updated.to_excel(sheet_path, index=False)
            return
        except:
            pass
    new_data.to_excel(sheet_path, index=False)

def consolidate_data(df: pd.DataFrame) -> pd.DataFrame:
    """Process raw data into invoice-ready format with original calculation logic"""
    consolidated = []
    for customer, group in df.groupby("MARK"):
        # Use session state defaults if they exist
        per_charge = (
            st.session_state.global_defaults['PER_CHARGES'] 
            if st.session_state.global_defaults.get('applied', False)
            else float(group["PER CHARGES"].iloc[0])
        )
        parking_charges = (
            st.session_state.global_defaults['PARKING_CHARGES']
            if st.session_state.global_defaults.get('applied', False)
            else float(group["PARKING CHARGES"].iloc[0])
        )
        weight_rate = (
            st.session_state.global_defaults['WEIGHT_RATE']
            if st.session_state.global_defaults.get('applied', False)
            else float(group["Weight Rate"].iloc[0]) if "Weight Rate" in group.columns else 1.0
        )
        
                # Calculate CBM for each row (weight-based vs. measured)
        group["Weight CBM"] = group["WEIGHT(KG)"] / weight_rate
        group["Actual CBM"] = group[["MEAS.(CBM)", "Weight CBM"]].max(axis=1)
        total_cbm = group["Actual CBM"].sum()
        total_qty = group["QTY"].astype(float).sum()  # ✅ Add this line
        # Calculate charges
        if total_cbm < 0.05:
            calculated_charges = 10.00  # Flat rate
            rate_applied = 10.00
        else:
            calculated_charges = (group["Actual CBM"] * per_charge).sum()
            rate_applied = per_charge

        total_charges = calculated_charges + parking_charges
        total_charges = calculated_charges + parking_charges
        first_row = group.iloc[0]
        
        # Multi-line fields
        fields = ["RECEIPT NO.", "QTY", "DESCRIPTION", "CBM", "WEIGHT(KG)"]
        joined = {f: "\n".join(group[f].astype(str)) for f in fields}
        
        # Build record with all required fields including rate
        consolidated.append({
            **joined,
            "PARKING CHARGES": f"{parking_charges:.2f}",
            "PER CHARGES": f"{rate_applied:.2f}",
            "RATE": f"{rate_applied:.2f}",
            "Weight Rate": f"{weight_rate:.2f}",
            "TOTAL CHARGES": f"{total_charges:.2f}",
            "MARK": customer,
            "CONTACT NUMBER": str(first_row.get("CONTACT NUMBER", "")),
            "CARGO NUMBER": str(first_row.get("CARGO NUMBER", "")),
            "TRACKING NUMBER": str(first_row.get("TRACKING NUMBER", "")),
            "TERMS": str(first_row.get("TERMS", "")),
            "TOTAL QTY": f"{total_qty:.2f}",
            "TOTAL CBM": f"{total_cbm:.2f}",
            "TOTAL CHARGES_SUM": f"{total_charges:.2f}",
            "FLAT_RATE_APPLIED": "Yes" if total_cbm < 0.05 else "No",
            "CALCULATED_CHARGES": calculated_charges
        })
    return pd.DataFrame(consolidated)


def create_download_link(file_path: str, label: str) -> str:
    """Generate HTML download link"""
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(file_path)}">{label}</a>'
    
def display_customer_markdowns(df: pd.DataFrame):
    """Display formatted markdown for each customer with editing capability"""
    st.header("📋 Customer Summaries & Editing")

    edited_df = df.copy()

    for idx, customer in enumerate(edited_df["MARK"].unique()):
        with st.expander(f"📌 {customer}", expanded=False):
            customer_data = edited_df[edited_df["MARK"] == customer].iloc[0]

            # Safe float conversion
            per_charges = float(customer_data.get('PER CHARGES', 0))
            parking_charges = float(customer_data.get('PARKING CHARGES', 0))
            total_cbm = float(customer_data.get('NUM_TOTAL_CBM', 0))  # use numeric CBM
            weight_rate = float(customer_data.get('Weight Rate', 1))

            # Recalculate charges
            if total_cbm < 0.05:
                calculated_charges = 10.00
            else:
                calculated_charges = total_cbm * per_charges
            total_charges = calculated_charges + parking_charges

            st.markdown(f"""
            **Customer:** {customer}  
            **Contact:** {customer_data.get('CONTACT NUMBER', 'N/A')}  
            **Total Items:** {customer_data.get('TOTAL QTY', 'N/A')}  
            **Total CBM:** {total_cbm:.2f}  
            **Weight Rate:** {weight_rate:.2f}  
            **Total Charges:** ${total_charges:.2f}  
            **Tracking Number:** {customer_data.get('TRACKING NUMBER', 'N/A')}  
            **Terms:** {customer_data.get('TERMS', 'N/A')}  
            **Flat Rate Applied:** {"Yes" if total_cbm < 0.05 else "No"}
            """)

            # Editable fields
            cols = st.columns(4)
            with cols[0]:
                new_per_charges = st.number_input("Per Charges ($/CBM)", value=per_charges, min_value=0.0, step=0.1, key=f"per_charges_{idx}")
            with cols[1]:
                new_parking = st.number_input("Parking Charges ($)", value=parking_charges, min_value=0.0, step=0.1, key=f"parking_{idx}")
            with cols[2]:
                new_weight_rate = st.number_input("Weight Rate (kg/CBM)", value=weight_rate, min_value=0.1, step=0.1, key=f"weight_rate_{idx}")
            with cols[3]:
                new_terms = st.text_input("Terms", value=str(customer_data.get('TERMS', '')), key=f"terms_{idx}")

            cols2 = st.columns(2)
            with cols2[0]:
                new_tracking = st.text_input("Tracking Number", value=str(customer_data.get('TRACKING NUMBER', '')), key=f"tracking_{idx}")
            with cols2[1]:
                new_contact = st.text_input("Contact Number", value=str(customer_data.get('CONTACT NUMBER', '')), key=f"contact_{idx}")

            if st.button(f"💾 Save Changes for {customer}", key=f"save_{idx}"):
                mask = edited_df["MARK"] == customer

                # Update edited fields
                edited_df.loc[mask, "PER CHARGES"] = float(new_per_charges)
                edited_df.loc[mask, "PARKING CHARGES"] = float(new_parking)
                edited_df.loc[mask, "Weight Rate"] = float(new_weight_rate)
                edited_df.loc[mask, "TERMS"] = new_terms
                edited_df.loc[mask, "TRACKING NUMBER"] = new_tracking
                edited_df.loc[mask, "CONTACT NUMBER"] = new_contact

                try:
                    original_df = st.session_state.raw_df
                    customer_rows = original_df[original_df["MARK"] == customer].copy()

                    customer_rows["WEIGHT(KG)"] = customer_rows["WEIGHT(KG)"].astype(float)
                    customer_rows["MEAS.(CBM)"] = customer_rows["MEAS.(CBM)"].astype(float)

                    weight_cbm = customer_rows["WEIGHT(KG)"] / float(new_weight_rate)
                    actual_cbm = pd.concat([customer_rows["MEAS.(CBM)"], weight_cbm], axis=1).max(axis=1)
                    current_cbm = actual_cbm.sum()

                    edited_df.loc[mask, "NUM_TOTAL_CBM"] = current_cbm
                    edited_df.loc[mask, "TOTAL CBM"] = f"{current_cbm:.2f}"

                except Exception as e:
                    st.warning(f"CBM recalculation failed: {e}")
                    current_cbm = float(edited_df.loc[mask, "NUM_TOTAL_CBM"].values[0])

                # Charges update
                current_cbm = float(edited_df.loc[mask, "NUM_TOTAL_CBM"].values[0])
                if current_cbm < 0.05:
                    calculated_charges = 10.00
                else:
                    calculated_charges = current_cbm * float(new_per_charges)

                total_charges = calculated_charges + float(new_parking)

                edited_df.loc[mask, "TOTAL CHARGES_SUM"] = total_charges
                edited_df.loc[mask, "TOTAL CHARGES"] = f"{total_charges:.2f}"
                edited_df.loc[mask, "FLAT_RATE_APPLIED"] = "Yes" if current_cbm < 0.05 else "No"
                edited_df.loc[mask, "RATE"] = f"{10.00:.2f}" if current_cbm < 0.05 else f"{new_per_charges:.2f}"

                # Save changes
                st.session_state.consolidated_df = edited_df
                st.success(f"✅ Changes saved and recalculated for {customer}!")
                st.rerun()


def consolidate_data(df: pd.DataFrame) -> pd.DataFrame:
    """Process raw data into invoice-ready format with calculation logic"""
    consolidated = []
    for customer, group in df.groupby("MARK"):
        # Use session state defaults if they exist
        defaults = st.session_state.get("global_defaults", {})
        use_defaults = defaults.get("applied", False)

        per_charge = defaults.get("PER_CHARGES", float(group["PER CHARGES"].iloc[0])) if use_defaults else float(group["PER CHARGES"].iloc[0])
        parking_charges = defaults.get("PARKING_CHARGES", float(group["PARKING CHARGES"].iloc[0])) if use_defaults else float(group["PARKING CHARGES"].iloc[0])
        weight_rate = defaults.get("WEIGHT_RATE", float(group["Weight Rate"].iloc[0])) if use_defaults else float(group["Weight Rate"].iloc[0])

        # Calculate CBM
        group["Weight CBM"] = group["WEIGHT(KG)"].astype(float) / weight_rate
        group["Actual CBM"] = group[["MEAS.(CBM)", "Weight CBM"]].astype(float).max(axis=1)
        total_cbm = group["Actual CBM"].sum()
        total_qty = group["QTY"].astype(float).sum()

        if total_cbm < 0.05:
            calculated_charges = 10.00
            rate_applied = 10.00
        else:
            calculated_charges = (group["Actual CBM"] * per_charge).sum()
            rate_applied = per_charge

        total_charges = calculated_charges + parking_charges

        first_row = group.iloc[0]
        fields = ["RECEIPT NO.", "QTY", "DESCRIPTION", "CBM", "WEIGHT(KG)"]
        joined = {f: "\n".join(group[f].astype(str)) for f in fields}

        consolidated.append({
            **joined,
            "PARKING CHARGES": f"{parking_charges:.2f}",
            "PER CHARGES": f"{rate_applied:.2f}",
            "RATE": f"{rate_applied:.2f}",
            "Weight Rate": f"{weight_rate:.2f}",
            "TOTAL CHARGES": f"{total_charges:.2f}",
            "MARK": customer,
            "CONTACT NUMBER": str(first_row.get("CONTACT NUMBER", "")),
            "CARGO NUMBER": str(first_row.get("CARGO NUMBER", "")),
            "TRACKING NUMBER": str(first_row.get("TRACKING NUMBER", "")),
            "TERMS": str(first_row.get("TERMS", "")),
            "TOTAL QTY": f"{total_qty:.2f}",
            "TOTAL CBM": f"{total_cbm:.2f}",
            "TOTAL CHARGES_SUM": total_charges,
            "FLAT_RATE_APPLIED": "Yes" if total_cbm < 0.05 else "No",
            "CALCULATED_CHARGES": calculated_charges,
            # preserve numeric total CBM
            "NUM_TOTAL_CBM": total_cbm
        })
    return pd.DataFrame(consolidated)

def create_download_link(file_path: str, label: str) -> str:
    """Generate HTML download link"""
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(file_path)}">{label}</a>'

def display_customer_markdowns(df: pd.DataFrame):
    """Display formatted markdown for each customer with editing capability"""
    st.header("📋 Customer Summaries & Editing")

    edited_df = df.copy()

    for idx, customer in enumerate(edited_df["MARK"].unique()):
        with st.expander(f"📌 {customer}", expanded=False):
            customer_data = edited_df[edited_df["MARK"] == customer].iloc[0]

            # Safe float conversion
            per_charges = float(customer_data.get('PER CHARGES', 0))
            parking_charges = float(customer_data.get('PARKING CHARGES', 0))
            total_cbm = float(customer_data.get('NUM_TOTAL_CBM', 0))  # use numeric CBM
            weight_rate = float(customer_data.get('Weight Rate', 1))

            # Recalculate charges
            if total_cbm < 0.05:
                calculated_charges = 10.00
            else:
                calculated_charges = total_cbm * per_charges
            total_charges = calculated_charges + parking_charges

            st.markdown(f"""
            **Customer:** {customer}  
            **Contact:** {customer_data.get('CONTACT NUMBER', 'N/A')}  
            **Total Items:** {customer_data.get('TOTAL QTY', 'N/A')}  
            **Total CBM:** {total_cbm:.2f}  
            **Weight Rate:** {weight_rate:.2f}  
            **Total Charges:** ${total_charges:.2f}  
            **Tracking Number:** {customer_data.get('TRACKING NUMBER', 'N/A')}  
            **Terms:** {customer_data.get('TERMS', 'N/A')}  
            **Flat Rate Applied:** {"Yes" if total_cbm < 0.05 else "No"}
            """)

            # Editable fields
            cols = st.columns(4)
            with cols[0]:
                new_per_charges = st.number_input("Per Charges ($/CBM)", value=per_charges, min_value=0.0, step=0.1, key=f"per_charges_{idx}")
            with cols[1]:
                new_parking = st.number_input("Parking Charges ($)", value=parking_charges, min_value=0.0, step=0.1, key=f"parking_{idx}")
            with cols[2]:
                new_weight_rate = st.number_input("Weight Rate (kg/CBM)", value=weight_rate, min_value=0.1, step=0.1, key=f"weight_rate_{idx}")
            with cols[3]:
                new_terms = st.text_input("Terms", value=str(customer_data.get('TERMS', '')), key=f"terms_{idx}")

            cols2 = st.columns(2)
            with cols2[0]:
                new_tracking = st.text_input("Tracking Number", value=str(customer_data.get('TRACKING NUMBER', '')), key=f"tracking_{idx}")
            with cols2[1]:
                new_contact = st.text_input("Contact Number", value=str(customer_data.get('CONTACT NUMBER', '')), key=f"contact_{idx}")

            # Save button inside expander
            if st.button(f"💾 Save Changes for {customer}", key=f"save_{idx}"):
                mask = edited_df["MARK"] == customer

                # Update edited fields
                edited_df.loc[mask, "PER CHARGES"] = float(new_per_charges)
                edited_df.loc[mask, "PARKING CHARGES"] = float(new_parking)
                edited_df.loc[mask, "Weight Rate"] = float(new_weight_rate)
                edited_df.loc[mask, "TERMS"] = new_terms
                edited_df.loc[mask, "TRACKING NUMBER"] = new_tracking
                edited_df.loc[mask, "CONTACT NUMBER"] = new_contact

                # Recalculate NUM_TOTAL_CBM based on new weight rate
                try:
                    original_df = st.session_state.raw_df
                    customer_rows = original_df[original_df["MARK"] == customer].copy()

                    customer_rows["WEIGHT(KG)"] = customer_rows["WEIGHT(KG)"].astype(float)
                    customer_rows["MEAS.(CBM)"] = customer_rows["MEAS.(CBM)"].astype(float)

                    weight_cbm = customer_rows["WEIGHT(KG)"] / float(new_weight_rate)
                    actual_cbm = pd.concat([customer_rows["MEAS.(CBM)"], weight_cbm], axis=1).max(axis=1)
                    current_cbm = actual_cbm.sum()

                    edited_df.loc[mask, "NUM_TOTAL_CBM"] = current_cbm
                    edited_df.loc[mask, "TOTAL CBM"] = f"{current_cbm:.2f}"

                except Exception as e:
                    st.warning(f"CBM recalculation failed: {e}")
                    current_cbm = float(edited_df.loc[mask, "NUM_TOTAL_CBM"].values[0])

                current_cbm = float(edited_df.loc[mask, "NUM_TOTAL_CBM"].values[0])
                if current_cbm < 0.05:
                    calculated_charges = 10.00
                else:
                    calculated_charges = current_cbm * float(new_per_charges)

                total_charges = calculated_charges + float(new_parking)

                edited_df.loc[mask, "TOTAL CHARGES_SUM"] = total_charges
                edited_df.loc[mask, "TOTAL CHARGES"] = f"{total_charges:.2f}"
                edited_df.loc[mask, "FLAT_RATE_APPLIED"] = "Yes" if current_cbm < 0.05 else "No"
                edited_df.loc[mask, "RATE"] = f"{10.00:.2f}" if current_cbm < 0.05 else f"{new_per_charges:.2f}"

                st.session_state.consolidated_df = edited_df
                st.success(f"Changes saved for {customer}!")
                st.rerun()

    return edited_df


def display_customer_editor():
    st.header("✏️ Edit Invoice Table Directly")

    # Load the current consolidated data
    df = st.session_state.consolidated_df.copy()
    
    st.subheader("📥 Download Packing List Print")

    # ✅ Packing list download
    if st.button("Create & Download Excel"):
        xls_buffer = export_to_excel_with_totals(st.session_state.consolidated_df)
        st.download_button(
            label="⬇️ Download Packing_List_Print.xlsx",
            data=xls_buffer,
            file_name="Packing_List_Print.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    # Display editable table (disable some columns)
    edited_df = st.data_editor(
        df,
        use_container_width=True,
        num_rows="dynamic",
        disabled=[
            "MARK", "TOTAL CBM", "TOTAL QTY", "FLAT_RATE_APPLIED",
            "TOTAL CHARGES", "TOTAL CHARGES_SUM", "RATE", "CALCULATED_CHARGES"
        ]
    )

    if st.button("💾 Save Table Changes"):
        # Pull raw data for accurate recalculation
        raw_df = st.session_state.raw_df.copy()

        # For each customer, update the original data with new editable values
        updated_records = []
        for _, row in edited_df.iterrows():
            customer = row["MARK"]
            per_charges = float(row["PER CHARGES"])
            parking = float(row["PARKING CHARGES"])
            weight_rate = float(row["Weight Rate"])

            # Get customer's raw rows
            customer_rows = raw_df[raw_df["MARK"] == customer].copy()
            customer_rows["WEIGHT(KG)"] = customer_rows["WEIGHT(KG)"].astype(float)
            customer_rows["MEAS.(CBM)"] = customer_rows["MEAS.(CBM)"].astype(float)
            customer_rows["QTY"] = customer_rows["QTY"].astype(float)

            # Recalculate CBM
            weight_cbm = customer_rows["WEIGHT(KG)"] / weight_rate
            actual_cbm = pd.concat([customer_rows["MEAS.(CBM)"], weight_cbm], axis=1).max(axis=1)
            total_cbm = actual_cbm.sum()
            total_qty = customer_rows["QTY"].sum()

            # Charges
            if total_cbm < 0.05:
                calculated_charges = 10.00
                rate_applied = 10.00
            else:
                calculated_charges = (actual_cbm * per_charges).sum()
                rate_applied = per_charges

            total_charges = calculated_charges + parking

            # Multi-line fields
            fields = ["RECEIPT NO.", "QTY", "DESCRIPTION", "CBM", "WEIGHT(KG)"]
            joined = {}
            for f in fields:
                if f == "CBM":
                    joined[f] = "\n".join(actual_cbm.round(3).astype(str))
                elif f in customer_rows.columns:
                    joined[f] = "\n".join(customer_rows[f].astype(str))
                else:
                    joined[f] = ""

            updated_records.append({
                **joined,
                "PARKING CHARGES": f"{parking:.2f}",
                "PER CHARGES": f"{rate_applied:.2f}",
                "RATE": f"{rate_applied:.2f}",
                "Weight Rate": f"{weight_rate:.2f}",
                "TOTAL CHARGES": f"{total_charges:.2f}",
                "MARK": customer,
                "CONTACT NUMBER": str(row.get("CONTACT NUMBER", "")),
                "CARGO NUMBER": str(row.get("CARGO NUMBER", "")),
                "TRACKING NUMBER": str(row.get("TRACKING NUMBER", "")),
                "TERMS": str(row.get("TERMS", "")),
                "TOTAL QTY": f"{total_qty:.2f}",
                "TOTAL CBM": f"{total_cbm:.2f}",
                "NUM_TOTAL_CBM": total_cbm,
                "TOTAL CHARGES_SUM": total_charges,
                "FLAT_RATE_APPLIED": "Yes" if total_cbm < 0.05 else "No",
                "CALCULATED_CHARGES": calculated_charges
            })

        # Save back to session
        st.session_state.consolidated_df = pd.DataFrame(updated_records)
        st.success("✅ Table updated and recalculated successfully.")
        st.rerun()


def main():
    st.title("📄 Invoice Generation System")
    
    # Initialize session state
    if 'global_defaults' not in st.session_state:
        st.session_state.global_defaults = {
            'PER_CHARGES': None,
            'WEIGHT_RATE': None,
            'PARKING_CHARGES': None,
            'applied': False
        }
    if 'consolidated_df' not in st.session_state:
        st.session_state.consolidated_df = None
    
    # Sidebar functions
    st.sidebar.header("⚙️ Quick Actions")
    st.sidebar.markdown("[📝 Sample Excel Template](#)")
    
    if st.sidebar.button("🔄 Clear All Data"):
        st.session_state.consolidated_df = None
        st.session_state.global_defaults = {
            'PER_CHARGES': None,
            'WEIGHT_RATE': None,
            'PARKING_CHARGES': None,
            'applied': False
        }
        if os.path.exists(OUTPUT_FOLDER):
            shutil.rmtree(OUTPUT_FOLDER)
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        st.rerun()
    
    if st.sidebar.button("ℹ️ Help"):
        st.sidebar.info("""
        **Instructions:**
        1. Upload Excel file with customer data
        2. Set global defaults if needed
        3. Generate individual or all invoices
        4. Download the generated PDFs
        """)

    # File Upload
    uploaded_file = st.file_uploader("📤 Upload Excel File", type=["xlsx", "xls"])
    if uploaded_file:
        try:
            df = pd.read_excel(uploaded_file)
            
            st.session_state.raw_df = df.copy()

            # Initialize missing columns with session state defaults if they exist
            per_charge_default = (
                st.session_state.global_defaults['PER_CHARGES'] 
                if st.session_state.global_defaults.get('applied', False)
                else 0.0
            )
            parking_default = (
                st.session_state.global_defaults['PARKING_CHARGES']
                if st.session_state.global_defaults.get('applied', False)
                else 0.0
            )
            weight_rate_default = (
                st.session_state.global_defaults['WEIGHT_RATE']
                if st.session_state.global_defaults.get('applied', False)
                else 1.0
            )
            
            # Initialize missing columns
            for col in ["PARKING CHARGES", "PER CHARGES", "Weight Rate", 
                       "TRACKING NUMBER", "TERMS"]:
                if col not in df.columns:
                    default_value = 0.0 if col in ["PARKING CHARGES", "PER CHARGES"] else (1.0 if col == "Weight Rate" else "")
                    df[col] = default_value
            
            # Handle zero weight rates
            if "Weight Rate" in df.columns:
                df["Weight Rate"] = df["Weight Rate"].replace(0, 1.0)
                
        except Exception as e:
            st.error(f"Error processing file: {e}")
            
        # Global Settings UI
        st.header("⚙️ Global Settings")
        col1, col2, col3 = st.columns(3)
        with col1:
            current_per_charge = float(df["PER CHARGES"].iloc[0])
            default_per_charge = st.number_input(
                "Default Per Charge ($/CBM)",
                value=st.session_state.global_defaults['PER_CHARGES'] or current_per_charge,
                min_value=0.0,
                step=0.1,
                key="global_per_charge"
            )
        with col2:
            current_weight_rate = float(df["Weight Rate"].iloc[0])
            default_weight_rate = st.number_input(
                "Default Weight Rate (kg/CBM)",
                value=st.session_state.global_defaults['WEIGHT_RATE'] or current_weight_rate,
                min_value=0.1,
                step=0.1,
                key="global_weight_rate"
            )
        with col3:
            current_parking = float(df["PARKING CHARGES"].iloc[0])
            default_parking = st.number_input(
                "Default Parking Charge ($)",
                value=st.session_state.global_defaults['PARKING_CHARGES'] or current_parking,
                min_value=0.0,
                step=0.1,
                key="global_parking"
            )
        
        if st.button("💾 Apply Global Settings", key="apply_global"):
            st.session_state.global_defaults = {
                'PER_CHARGES': default_per_charge,
                'WEIGHT_RATE': default_weight_rate,
                'PARKING_CHARGES': default_parking,
                'applied': True
            }
            st.success("Global settings applied to all customers!")
        
        # Process data with original calculation logic
        df["Weight CBM"] = df["WEIGHT(KG)"] / df["Weight Rate"]
        df["CBM"] = df[["MEAS.(CBM)", "Weight CBM"]].max(axis=1)
        st.session_state.consolidated_df = consolidate_data(df)
        
        # Choose editing method
        st.subheader("✏️ Customer Editing Mode")
        use_table_editor = st.checkbox("Use Table Editor Instead of Per-Customer Forms", value=True)

        if use_table_editor:
            display_customer_editor()
        else:
            display_customer_markdowns(st.session_state.consolidated_df)
        
        # Display processed table
        st.header("📊 Processed Data")
        st.dataframe(st.session_state.consolidated_df)
        
        # Invoice Generation
        st.header("🖨️ Invoice Generation")
        start_num = st.number_input(
            "Starting Invoice Number", 
            min_value=1, 
            value=1,
            key="invoice_start"
        )
        # Exclusion option before generating all invoices
        st.subheader("🙅‍♂️ Exclude Customers from Invoice Generation")
        all_customers = st.session_state.consolidated_df["MARK"].unique().tolist()
        excluded_customers = st.multiselect("Select customers to exclude from 'Generate All Invoices'", options=all_customers, key="exclude_customers")

        
        
        if st.button("🔄 Generate All Invoices", key="generate_all"):
            if os.path.exists(OUTPUT_FOLDER):
                shutil.rmtree(OUTPUT_FOLDER)
            os.makedirs(OUTPUT_FOLDER)
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
        
            included_df = st.session_state.consolidated_df[~st.session_state.consolidated_df["MARK"].isin(excluded_customers)].reset_index(drop=True)
            for i, (_, row) in enumerate(included_df.iterrows()):
    
                status_text.text(f"Processing {i+1}/{len(st.session_state.consolidated_df)}: {row['MARK']}")
                progress_bar.progress((i + 1) / len(st.session_state.consolidated_df))
                
                pdf_path = generate_pdf_from_template(
                    TEMPLATE_PATH,
                    row.to_dict(),
                    OUTPUT_FOLDER,
                    start_num + i
                )
                
                if pdf_path:
                    update_notification_sheet(
                        OUTPUT_FOLDER,
                        os.path.basename(pdf_path),
                        row["MARK"],
                        start_num + i,
                        row["CONTACT NUMBER"],
                        row["TOTAL CHARGES_SUM"]
                    )
            
            # Create ZIP archive
            zip_path = os.path.join(OUTPUT_FOLDER, "invoices.zip")
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in os.listdir(OUTPUT_FOLDER):
                    if file.endswith('.pdf'):
                        zipf.write(
                            os.path.join(OUTPUT_FOLDER, file), 
                            file
                        )
            
            st.markdown(
                create_download_link(zip_path, "📥 Download All Invoices"),
                unsafe_allow_html=True
            )
            st.success("✅ All invoices generated successfully!")
        
        # Single Invoice Generation
        st.subheader("🖨️ Single Invoice")
        if st.session_state.consolidated_df is not None:
            customer = st.selectbox(
                "Select Customer",
                options=st.session_state.consolidated_df["MARK"].unique(),
                key="customer_select"
            )
            single_num = st.number_input(
                "Invoice Number", 
                min_value=1, 
                value=start_num,
                key="single_invoice_num"
            )
            
            if st.button("🖨️ Generate Selected Invoice", key="generate_single"):
                row = st.session_state.consolidated_df[
                    st.session_state.consolidated_df["MARK"] == customer
                ].iloc[0]
                
                pdf_path = generate_pdf_from_template(
                    TEMPLATE_PATH,
                    row.to_dict(),
                    OUTPUT_FOLDER,
                    single_num
                )
                
                if pdf_path:
                    with open(pdf_path, "rb") as f:
                        st.download_button(
                            "📥 Download Invoice",
                            f,
                            file_name=os.path.basename(pdf_path),
                            mime="application/pdf"
                        )
                    st.success(f"✅ Invoice #{single_num} generated for {customer}!")

if __name__ == "__main__":
    main()

